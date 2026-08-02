import io
from docx import Document


def extract_text_from_docx(content: bytes) -> dict:
    try:
        document = Document(io.BytesIO(content))
    except Exception:
        return {"success": False, "error": "The DOCX file appears to be corrupted and could not be opened."}

    parts = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    full_text = "\n".join(parts)

    if not full_text.strip():
        return {"success": False, "error": "No readable text found in this document."}

    return {"success": True, "text": full_text}